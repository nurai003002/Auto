import io
import json
from datetime import date

from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_http_methods
from django.views.decorators.csrf import csrf_exempt

import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment
from docx import Document
from docx.shared import Pt, RGBColor

from apps.cars.models import Car, Fix


# ── Index ─────────────────────────────────────────────────────

def index(request):
    today = date.today()
    cars = list(Car.objects.all())
    fixes = list(Fix.objects.select_related('car').order_by('-date'))

    # Group fixes by car
    fixes_by_car = {}
    for f in fixes:
        fixes_by_car.setdefault(f.car_id, []).append(f)

    def repair_status(next_date):
        if not next_date:
            return 'ok', 'В норме'
        diff = (next_date - today).days
        if diff < 0:
            return 'overdue', 'Просрочено'
        if diff <= 31:
            return 'soon', 'Скоро'
        return 'ok', 'В норме'

    cars_data = []
    for car in cars:
        car_fixes = fixes_by_car.get(car.id, [])
        last_fix = car_fixes[0] if car_fixes else None
        upcoming = [f for f in car_fixes if f.next_date]
        next_fix = min(upcoming, key=lambda f: f.next_date) if upcoming else None
        status, status_label = repair_status(next_fix.next_date if next_fix else None)
        cars_data.append({
            'car': car,
            'last_fix': last_fix,
            'next_fix': next_fix,
            'status': status,
            'status_label': status_label,
        })

    # Stats
    total_cars = len(cars)
    total_fixes = len(fixes)
    overdue_count = sum(1 for c in cars_data if c['status'] == 'overdue')
    soon_count = sum(1 for c in cars_data if c['status'] == 'soon')

    # Reminders (overdue, this month, next month)
    this_month = today.month
    this_year = today.year
    next_month = (this_month % 12) + 1
    next_month_year = this_year + (1 if this_month == 12 else 0)

    reminders = []
    for f in fixes:
        if not f.next_date:
            continue
        diff = (f.next_date - today).days
        nd = f.next_date
        if diff < 0:
            status = 'overdue'
            label = f'Просрочено на {abs(diff)} дн.'
        elif nd.month == this_month and nd.year == this_year:
            status = 'this_month'
            label = 'Сегодня!' if diff == 0 else f'Через {diff} дн.'
        elif nd.month == next_month and nd.year == next_month_year:
            status = 'next_month'
            label = f'Через {diff} дн.'
        else:
            continue
        reminders.append({
            'id': f.id,
            'car': f'{f.car.marka} {f.car.model}',
            'plate': f.car.gos_munber,
            'type': f.repair_type,
            'date': f.next_date,
            'diff': diff,
            'status': status,
            'label': label,
        })
    reminders.sort(key=lambda r: r['diff'])

    alert_reminder_ids = [
        r['id'] for r in reminders if r['status'] in ('overdue', 'this_month')
    ]

    context = {
        'cars_data': cars_data,
        'cars': cars,
        'fixes': fixes,
        'reminders': reminders,
        'total_cars': total_cars,
        'total_fixes': total_fixes,
        'overdue_count': overdue_count,
        'soon_count': soon_count,
        'alert_count': overdue_count + soon_count,
        'recent_cars_data': list(reversed(cars_data))[:5],
        'alert_reminder_ids_json': json.dumps(alert_reminder_ids),
    }
    return render(request, 'index.html', context)


def car_to_dict(car):
    return {
        'id': car.id,
        'brand': car.marka,
        'model': car.model,
        'plate': car.gos_munber,
        'year': car.year,
        'note': car.note,
    }


@csrf_exempt
@require_http_methods(['GET', 'POST'])
def cars_list(request):
    if request.method == 'GET':
        cars = list(Car.objects.all().values('id', 'marka', 'model', 'gos_munber', 'year', 'note'))
        data = [{'id': c['id'], 'brand': c['marka'], 'model': c['model'],
                 'plate': c['gos_munber'], 'year': c['year'], 'note': c['note']} for c in cars]
        return JsonResponse(data, safe=False)

    body = json.loads(request.body)
    car = Car.objects.create(
        marka=body['brand'],
        model=body['model'],
        gos_munber=body['plate'],
        year=body['year'],
        note=body.get('note', ''),
    )
    return JsonResponse(car_to_dict(car), status=201)


@csrf_exempt
@require_http_methods(['GET', 'PUT', 'DELETE'])
def cars_detail(request, pk):
    try:
        car = Car.objects.get(pk=pk)
    except Car.DoesNotExist:
        return JsonResponse({'error': 'Not found'}, status=404)

    if request.method == 'GET':
        return JsonResponse(car_to_dict(car))

    if request.method == 'PUT':
        body = json.loads(request.body)
        car.marka = body.get('brand', car.marka)
        car.model = body.get('model', car.model)
        car.gos_munber = body.get('plate', car.gos_munber)
        car.year = body.get('year', car.year)
        car.note = body.get('note', car.note)
        car.save()
        return JsonResponse(car_to_dict(car))

    car.delete()
    return JsonResponse({}, status=204)


def fix_to_dict(fix):
    return {
        'id': fix.id,
        'car_id': fix.car_id,
        'type': fix.repair_type,
        'date': str(fix.date),
        'next_date': str(fix.next_date) if fix.next_date else None,
        'comment': fix.description,
    }


@csrf_exempt
@require_http_methods(['GET', 'POST'])
def repairs_list(request):
    if request.method == 'GET':
        fixes = Fix.objects.all().order_by('-date')
        return JsonResponse([fix_to_dict(f) for f in fixes], safe=False)

    body = json.loads(request.body)
    fix = Fix.objects.create(
        car_id=body['car_id'],
        repair_type=body['type'],
        date=body['date'],
        next_date=body.get('next_date') or None,
        description=body.get('comment', ''),
    )
    return JsonResponse(fix_to_dict(fix), status=201)


@csrf_exempt
@require_http_methods(['PUT', 'DELETE'])
def repairs_detail(request, pk):
    try:
        fix = Fix.objects.get(pk=pk)
    except Fix.DoesNotExist:
        return JsonResponse({'error': 'Not found'}, status=404)

    if request.method == 'PUT':
        body = json.loads(request.body)
        fix.car_id = body.get('car_id', fix.car_id)
        fix.repair_type = body.get('type', fix.repair_type)
        fix.date = body.get('date', fix.date)
        fix.next_date = body.get('next_date') or None
        fix.description = body.get('comment', fix.description)
        fix.save()
        return JsonResponse(fix_to_dict(fix))

    fix.delete()
    return JsonResponse({}, status=204)

def _get_report_data(report_type):
    today = date.today()

    if report_type == 'cars':
        headers = ['Марка', 'Модель', 'Гос. номер', 'Год', 'Примечание']
        rows = [
            [c.marka, c.model, c.gos_munber, c.year, c.note]
            for c in Car.objects.all().order_by('marka', 'model')
        ]
        return headers, rows

    if report_type == 'repairs':
        headers = ['Автомобиль', 'Тип ремонта', 'Дата ремонта', 'Следующий ремонт', 'Комментарий']
        rows = []
        for f in Fix.objects.select_related('car').order_by('-date'):
            rows.append([
                f'{f.car.marka} {f.car.model} ({f.car.gos_munber})',
                f.repair_type,
                f.date.strftime('%d.%m.%Y'),
                f.next_date.strftime('%d.%m.%Y') if f.next_date else '—',
                f.description,
            ])
        return headers, rows

    
    headers = ['Автомобиль', 'Гос. номер', 'Тип ремонта', 'Дата следующего ремонта', 'Статус']
    rows = []
    for f in Fix.objects.select_related('car').exclude(next_date=None).order_by('next_date'):
        diff = (f.next_date - today).days
        if diff < 0:
            status = 'Просрочено'
        elif diff <= 31:
            status = 'Скоро'
        else:
            continue
        rows.append([
            f'{f.car.marka} {f.car.model}',
            f.car.gos_munber,
            f.repair_type,
            f.next_date.strftime('%d.%m.%Y'),
            status,
        ])
    return headers, rows


def export_excel(request, report_type):
    headers, rows = _get_report_data(report_type)
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Отчёт'

    header_fill = PatternFill('solid', fgColor='1E3A5F')
    header_font = Font(bold=True, color='FFFFFF', size=11)
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')

    for row_idx, row in enumerate(rows, 2):
        for col_idx, value in enumerate(row, 1):
            ws.cell(row=row_idx, column=col_idx, value=value)

    for col in ws.columns:
        max_len = max((len(str(cell.value or '')) for cell in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    response = HttpResponse(
        buf,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = f'attachment; filename="autotrack_{report_type}.xlsx"'
    return response


def export_word(request, report_type):
    headers, rows = _get_report_data(report_type)
    titles = {
        'cars': 'Список автомобилей',
        'repairs': 'История ремонтов',
        'upcoming': 'Предстоящее обслуживание',
    }
    doc = Document()
    title = doc.add_heading(titles.get(report_type, 'Отчёт'), level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    doc.add_paragraph(f'Дата формирования: {date.today().strftime("%d.%m.%Y")}')
    doc.add_paragraph('')

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)

    for row_idx, row in enumerate(rows, 1):
        cells = table.rows[row_idx].cells
        for col_idx, value in enumerate(row):
            cells[col_idx].text = str(value)
            cells[col_idx].paragraphs[0].runs[0].font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    response = HttpResponse(
        buf,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = f'attachment; filename="autotrack_{report_type}.docx"'
    return response
