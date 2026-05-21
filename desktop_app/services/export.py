"""
AutoTrack — Export service (Excel, Word, PDF).
"""
import io
from datetime import date


def export_to_excel(path, title, headers, rows):
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Отчёт"

    header_fill = PatternFill('solid', fgColor='1E3A5F')
    header_font = Font(bold=True, color='FFFFFF', size=11)
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal='center', vertical='center')

    for row_idx, row in enumerate(rows, 2):
        for col_idx, value in enumerate(row, 1):
            ws.cell(row=row_idx, column=col_idx, value=value)

    for col in ws.columns:
        max_len = max((len(str(cell.value or '')) for cell in col), default=10)
        ws.column_dimensions[col[0].column_letter].width = min(max_len + 4, 50)

    wb.save(path)


def export_to_word(path, title, headers, rows):
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    doc.add_paragraph(f'Дата формирования: {date.today().strftime("%d.%m.%Y")}')
    doc.add_paragraph('')

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)

    for row_idx, row in enumerate(rows, 1):
        cells = table.rows[row_idx].cells
        for col_idx, value in enumerate(row):
            cells[col_idx].text = str(value)
            cells[col_idx].paragraphs[0].runs[0].font.size = Pt(10)

    doc.save(path)


def export_to_pdf(path, title, headers, rows):
    try:
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib import colors
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont

        doc = SimpleDocTemplate(path, pagesize=landscape(A4))
        elements = []
        styles = getSampleStyleSheet()

        elements.append(Paragraph(title, styles['Title']))
        elements.append(Spacer(1, 12))
        elements.append(Paragraph(
            f'Дата: {date.today().strftime("%d.%m.%Y")}', styles['Normal']
        ))
        elements.append(Spacer(1, 12))

        data = [headers] + rows
        table = Table(data)
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1E3A5F')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')]),
            ('TOPPADDING', (0, 0), (-1, -1), 4),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
        ]))
        elements.append(table)
        doc.build(elements)
    except ImportError:
        raise ImportError("Для PDF экспорта установите reportlab: pip install reportlab")
